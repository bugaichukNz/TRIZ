"""Генерация DOCX-отчёта TRIZ-анализа для руководства и внедрения."""

from __future__ import annotations

import io
import os
from datetime import datetime
from typing import Any

try:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
    from docx.table import _Cell
except ImportError as exc:
    raise ImportError(
        "Для DOCX установите зависимости: pip install python-docx matplotlib"
    ) from exc

def _total_score(sol: dict[str, Any]) -> float:
    return round(
        sol.get("effectiveness_score", 0)
        + sol.get("scalability_score", 0)
        - (sol.get("complexity_score", 0) + sol.get("cost_score", 0)) / 2,
        1,
    )

DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MID_TEXT = RGBColor(0x6B, 0x6B, 0x6B)
HEADER_BG = "F0F0F0"
ACCENT = "#10A37F"
FONT_NAME = "Arial"


def _set_cell_shading(cell: _Cell, fill_hex: str) -> None:
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.get_or_add_shd()
    shd.set(qn("w:fill"), fill_hex)


def _style_header_row(row, cols: int) -> None:
    for i in range(cols):
        cell = row.cells[i]
        _set_cell_shading(cell, HEADER_BG)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = FONT_NAME
                run.font.color.rgb = DARK_TEXT
                run.font.size = Pt(10)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = DARK_TEXT if level == 1 else MID_TEXT


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(11)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)


def _build_comparison_chart(solutions: list[dict[str, Any]]) -> io.BytesIO:
    """Bar chart: эффективность vs сложность внедрения."""
    labels = [s.get("title", f"#{s.get('id', '?')}")[:28] for s in solutions]
    effectiveness = [s.get("effectiveness_score", 0) for s in solutions]
    complexity = [s.get("complexity_score", 0) for s in solutions]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = range(len(labels))
    width = 0.35
    ax.bar(
        [i - width / 2 for i in x],
        effectiveness,
        width,
        label="Эффективность",
        color=ACCENT,
    )
    ax.bar(
        [i + width / 2 for i in x],
        complexity,
        width,
        label="Сложность внедрения",
        color="#404040",
        alpha=0.85,
    )
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=15, ha="right", fontsize=9)
    ax.set_ylim(0, 10)
    ax.set_ylabel("Оценка (1–10)")
    ax.set_title("Сравнение решений: эффективность vs сложность", fontsize=11)
    ax.legend(loc="upper right", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()

    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def build_report_docx(
    problem: str,
    result: dict[str, Any],
    *,
    generated_at: datetime | None = None,
    analyst_name: str | None = None,
    version: str = "1.0",
) -> bytes:
    """Собирает DOCX-отчёт и возвращает байты файла."""
    ts = generated_at or datetime.now()
    analyst = analyst_name or os.getenv("REPORT_ANALYST_NAME", "TRIZ AI-Ассистент")

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # --- Титульная страница ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TRIZ-аналитический отчёт")
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = FONT_NAME
    run.font.color.rgb = DARK_TEXT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(problem[:200] + ("…" if len(problem) > 200 else ""))
    r.font.size = Pt(12)
    r.font.name = FONT_NAME
    r.font.color.rgb = MID_TEXT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        f"Дата: {ts.strftime('%d.%m.%Y %H:%M')}",
        f"Аналитик: {analyst}",
        f"Версия: {version}",
    ):
        mr = meta.add_run(line + "\n")
        mr.font.size = Pt(10)
        mr.font.name = FONT_NAME

    doc.add_page_break()

    # --- 1. Описание задачи ---
    _add_heading(doc, "1. Описание задачи")
    _add_body(doc, result.get("problem_description", problem))
    assumptions = result.get("assumptions") or []
    if assumptions:
        _add_heading(doc, "Принятые допущения", level=2)
        _add_bullets(doc, assumptions)

    # --- 2. Система и надсистема ---
    _add_heading(doc, "2. Система и надсистема")
    ctx = result.get("system_context") or {}
    sys_table = doc.add_table(rows=1, cols=2)
    sys_table.style = "Table Grid"
    hdr = sys_table.rows[0].cells
    hdr[0].text = "Элемент"
    hdr[1].text = "Описание"
    _style_header_row(sys_table.rows[0], 2)
    rows_data = [
        ("Система", ctx.get("system", "—")),
        ("Надсистема", ctx.get("supersystem", "—")),
        ("Подсистемы", "; ".join(ctx.get("subsystems") or []) or "—"),
        ("Полезные функции", "; ".join(ctx.get("useful_functions") or []) or "—"),
        ("Нежелательные эффекты", "; ".join(ctx.get("harmful_effects") or []) or "—"),
        ("Ограничения", "; ".join(ctx.get("constraints") or []) or "—"),
        ("Доступные ресурсы", "; ".join(ctx.get("resources") or []) or "—"),
    ]
    for label, value in rows_data:
        row = sys_table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(10)

    # --- 3. Противоречия и ИКР ---
    _add_heading(doc, "3. Противоречия и ИКР")
    contra_table = doc.add_table(rows=1, cols=2)
    contra_table.style = "Table Grid"
    ch = contra_table.rows[0].cells
    ch[0].text = "Тип"
    ch[1].text = "Формулировка"
    _style_header_row(contra_table.rows[0], 2)
    for label, key in (
        ("Техническое противоречие (ТП)", "technical_contradiction"),
        ("Физическое противоречие (ФП)", "physical_contradiction"),
        ("Тип", "contradiction_type"),
    ):
        r = contra_table.add_row().cells
        r[0].text = label
        r[1].text = str(result.get(key, "—"))

    ifr_p = doc.add_paragraph()
    ifr_run = ifr_p.add_run(f"ИКР: {result.get('ideal_final_result', '—')}")
    ifr_run.bold = True
    ifr_run.font.name = FONT_NAME
    ifr_run.font.color.rgb = DARK_TEXT
    ifr_run.font.size = Pt(11)

    # --- 4. Анализ (кратко) ---
    _add_heading(doc, "4. Анализ")
    analysis = result.get("analysis") or {}
    for title, key in (
        ("Причинно-следственные цепочки", "causal_chains"),
        ("Функциональный анализ", "functional_analysis"),
        ("Выявленные ресурсы", "resources_analysis"),
        ("Ключевые зоны противоречий", "contradiction_zones"),
    ):
        _add_heading(doc, title, level=2)
        _add_body(doc, str(analysis.get(key, "—")))

    # --- 5. Инструменты ТРИЗ ---
    _add_heading(doc, "5. Применённые инструменты ТРИЗ")
    tools = result.get("triz_tools") or []
    t_table = doc.add_table(rows=1, cols=4)
    t_table.style = "Table Grid"
    th = t_table.rows[0].cells
    for i, h in enumerate(
        ("Инструмент", "Почему применён", "Что выявил", "Практическая ценность")
    ):
        th[i].text = h
    _style_header_row(t_table.rows[0], 4)
    for tool in tools:
        row = t_table.add_row().cells
        row[0].text = tool.get("tool", "—")
        row[1].text = tool.get("why_applied", "—")
        row[2].text = tool.get("insight", "—")
        row[3].text = tool.get("practical_value", "—")

    # --- 6. Решения ---
    _add_heading(doc, "6. Решения")
    concepts = result.get("solution_concepts") or []
    s_table = doc.add_table(rows=1, cols=5)
    s_table.style = "Table Grid"
    sh = s_table.rows[0].cells
    for i, h in enumerate(
        ("Решение", "Принцип ТРИЗ", "Механизм", "Применимость", "Риски / ограничения")
    ):
        sh[i].text = h
    _style_header_row(s_table.rows[0], 5)
    for sol in concepts:
        row = s_table.add_row().cells
        row[0].text = sol.get("title", "—")
        row[1].text = sol.get("triz_principle", "—")
        row[2].text = sol.get("mechanism", "—")
        row[3].text = sol.get("applicability", "—")
        row[4].text = sol.get("risks", "—")

    # --- 7. Сравнение решений ---
    _add_heading(doc, "7. Сравнение решений")
    c_table = doc.add_table(rows=1, cols=6)
    c_table.style = "Table Grid"
    chdr = c_table.rows[0].cells
    for i, h in enumerate(
        (
            "Решение",
            "Эффективность",
            "Сложность",
            "Стоимость",
            "Масштабируемость",
            "Итоговый балл",
        )
    ):
        chdr[i].text = h
    _style_header_row(c_table.rows[0], 6)
    for sol in concepts:
        row = c_table.add_row().cells
        row[0].text = sol.get("title", "—")
        row[1].text = str(sol.get("effectiveness_score", "—"))
        row[2].text = str(sol.get("complexity_score", "—"))
        row[3].text = str(sol.get("cost_score", "—"))
        row[4].text = str(sol.get("scalability_score", "—"))
        row[5].text = str(_total_score(sol))

    # --- 8. Визуализация ---
    _add_heading(doc, "8. Визуализация")
    if concepts:
        chart_buf = _build_comparison_chart(concepts)
        doc.add_picture(chart_buf, width=Inches(6.2))

    # --- 9. Рекомендации ---
    _add_heading(doc, "9. Рекомендации к внедрению")
    rec = result.get("recommendations") or {}
    if rec.get("priorities"):
        _add_heading(doc, "Приоритеты", level=2)
        _add_bullets(doc, rec["priorities"])
    for title, key in (
        ("Быстрые проверки гипотез", "quick_checks"),
        ("MVP / пилоты", "mvp_pilots"),
        ("Критические риски", "critical_risks"),
        ("Требуемые эксперименты", "experiments"),
        ("Метрики эффективности", "metrics"),
    ):
        items = rec.get(key) or []
        if items:
            _add_heading(doc, title, level=2)
            _add_bullets(doc, items)

    # --- 10. Итоговый вывод ---
    _add_heading(doc, "10. Итоговый вывод")
    conclusion = result.get("final_conclusion") or {}
    sum_table = doc.add_table(rows=1, cols=2)
    sum_table.style = "Table Grid"
    sum_table.rows[0].cells[0].text = "Параметр"
    sum_table.rows[0].cells[1].text = "Значение"
    _style_header_row(sum_table.rows[0], 2)
    for label, key in (
        ("Рекомендуемое решение", "recommended_solution"),
        ("Ключевой риск", "key_risk"),
        ("Следующий шаг", "next_step"),
    ):
        r = sum_table.add_row().cells
        r[0].text = label
        r[1].text = str(conclusion.get(key, "—"))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
